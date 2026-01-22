from docx import Document

doc = Document()
doc.add_heading('Lumina Test Document', 0)
doc.add_paragraph('Hello, {{ name }}!')
doc.add_paragraph('Today is {{ date }}.')
doc.save('test_template.docx')
print("Created test_template.docx")
